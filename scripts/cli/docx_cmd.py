from __future__ import annotations

import json
import re
import zipfile
from xml.etree import ElementTree
from pathlib import Path
from typing import Any, Dict, List, Optional

from core.search import resolve_crossref
from cli._common import (
    _output, _write_docx_from_markdown,
    _get_or_create_footnotes_part, _add_footnote_to_element,
)


def _para_replace_text(paragraph, find_text, replace_text):
    """段落级文本替换：先尝试单 run 内替换；跨 run 时精确切割，只改被命中区间，保留其余 run 格式。"""
    for run in paragraph.runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text, 1)
            return True

    runs = paragraph.runs
    if not runs:
        return False

    boundaries = []
    pos = 0
    for run in runs:
        end = pos + len(run.text)
        boundaries.append((pos, end, run))
        pos = end

    full = "".join(r.text for r in runs)
    if find_text not in full:
        return False

    idx = full.index(find_text)
    end_idx = idx + len(find_text)

    first_i = last_i = None
    for i, (s, e, _) in enumerate(boundaries):
        if first_i is None and s <= idx < e:
            first_i = i
        if s < end_idx <= e:
            last_i = i
            break

    if first_i is None or last_i is None:
        return False

    fs, fe, first_run = boundaries[first_i]
    ls, le, last_run = boundaries[last_i]

    if first_i == last_i:
        offset = idx - fs
        first_run.text = first_run.text[:offset] + replace_text + first_run.text[offset + len(find_text):]
    else:
        first_run.text = first_run.text[:idx - fs] + replace_text
        for j in range(first_i + 1, last_i):
            boundaries[j][2].text = ""
        last_run.text = last_run.text[end_idx - ls:]

    return True


def _find_run_containing(paragraph, text):
    """在段落中找到包含指定文本的 run（取全文第一次匹配），返回 run._element 或 None。
    跨 run 时精确定位到匹配文本末尾字符所在的 run（脚注应插在该 run 之后）。
    """
    if not text:
        return None

    for run in paragraph.runs:
        if text in run.text:
            return run._element

    runs = paragraph.runs
    if not runs:
        return None

    boundaries = []
    pos = 0
    for run in runs:
        end = pos + len(run.text)
        boundaries.append((pos, end, run))
        pos = end

    full = "".join(r.text for r in runs)
    if text not in full:
        return None

    end_idx = full.index(text) + len(text)
    target = end_idx - 1
    for start, end, run in boundaries:
        if start <= target < end:
            return run._element

    return runs[-1]._element


_WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W = f"{{{_WORD_NS}}}"


def _ooxml_paragraph_texts(element) -> List[str]:
    paragraphs = []
    for paragraph in element.iter(f"{_W}p"):
        text = "".join(
            node.text or "" for node in paragraph.iter(f"{_W}t")
        ).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def _extract_docx_auxiliary_text(filepath: Path) -> tuple:
    """Extract text held outside python-docx's body/table collections."""
    sections: List[str] = []
    observed = ["body", "tables"]
    with zipfile.ZipFile(filepath) as archive:
        names = set(archive.namelist())

        note_parts = (
            ("word/footnotes.xml", "footnotes", "脚注"),
            ("word/endnotes.xml", "endnotes", "尾注"),
            ("word/comments.xml", "comments", "批注"),
        )
        for part_name, part_id, heading in note_parts:
            if part_name not in names:
                continue
            root = ElementTree.fromstring(archive.read(part_name))
            paragraphs = []
            if part_id in {"footnotes", "endnotes"}:
                item_tag = f"{_W}{'footnote' if part_id == 'footnotes' else 'endnote'}"
                for item in root.iter(item_tag):
                    note_type = item.get(f"{_W}type")
                    note_id = item.get(f"{_W}id", "")
                    if note_type or note_id in {"-1", "0", "1"}:
                        continue
                    paragraphs.extend(_ooxml_paragraph_texts(item))
            else:
                paragraphs = _ooxml_paragraph_texts(root)
            if paragraphs:
                observed.append(part_id)
                sections.append(f"[{heading}]\n" + "\n".join(paragraphs))

        for prefix, part_id, heading in (
            ("word/header", "headers", "页眉"),
            ("word/footer", "footers", "页脚"),
        ):
            paragraphs = []
            for part_name in sorted(
                name for name in names
                if name.startswith(prefix) and name.endswith(".xml")
            ):
                root = ElementTree.fromstring(archive.read(part_name))
                paragraphs.extend(_ooxml_paragraph_texts(root))
            if paragraphs:
                observed.append(part_id)
                sections.append(f"[{heading}]\n" + "\n".join(paragraphs))

        if "word/document.xml" in names:
            root = ElementTree.fromstring(archive.read("word/document.xml"))
            textboxes = []
            for textbox in root.iter(f"{_W}txbxContent"):
                textboxes.extend(_ooxml_paragraph_texts(textbox))
            if textboxes:
                observed.append("textboxes")
                sections.append("[文本框]\n" + "\n".join(textboxes))

    observability = {
        "observed_parts": list(dict.fromkeys(observed)),
        "unobserved_parts": [
            "embedded_objects",
            "images_without_ocr",
            "drawing_semantics",
            "tracked_deletions",
        ],
    }
    return sections, observability


def cmd_read_paper(args):
    """读取用户论文文件（.docx / .txt / .md），输出 UTF-8 纯文本"""
    filepath = Path(args.filepath)
    if not filepath.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND", "message": f"文件不存在: {args.filepath}"})
        return

    suffix = filepath.suffix.lower()
    text = ""
    observability = {
        "observed_parts": ["plain_text"],
        "unobserved_parts": [],
    }

    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            _output({"status": "error", "code": "MISSING_DEPENDENCY",
                     "message": "缺少 python-docx 依赖"})
            return
        try:
            doc = Document(str(filepath))
            parts: List[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_texts:
                        parts.append(" | ".join(row_texts))
            auxiliary, observability = _extract_docx_auxiliary_text(filepath)
            parts.extend(auxiliary)
            text = "\n\n".join(parts)
        except Exception as e:
            _output({"status": "error", "code": "DOCX_PARSE_FAILED", "message": f"docx 解析失败: {e}"})
            return

    elif suffix in (".txt", ".md", ".markdown"):
        decoded = False
        for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030"):
            try:
                text = filepath.read_text(encoding=enc)
                decoded = True
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if not decoded:
            _output({"status": "error", "code": "ENCODING_ERROR", "message": "无法识别文件编码"})
            return

    elif suffix == ".pdf":
        _output({"status": "error", "code": "UNSUPPORTED_FORMAT",
                 "message": "PDF 请使用 Agent 内置的文件读取工具直接读取，无需 read-paper"})
        return
    else:
        _output({"status": "error", "code": "UNSUPPORTED_FORMAT",
                 "message": f"不支持的文件格式: {suffix}"})
        return

    char_count = len(text)
    para_count = text.count("\n\n") + 1 if text else 0

    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(text, encoding="utf-8")
        _output({"status": "success",
                 "message": f"已提取到: {args.output}",
                 "chars": char_count,
                 "paragraphs": para_count,
                 "observability": observability})
    else:
        if args.raw:
            print(text)
        else:
            _output({"status": "success",
                     "chars": char_count,
                     "paragraphs": para_count,
                     "observability": observability,
                     "text": text})


def cmd_pdf_meta(args):
    """从 PDF 文件中提取元数据（标题、作者、DOI 等）"""
    filepath = Path(args.filepath)
    if not filepath.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND", "message": f"文件不存在: {args.filepath}"})
        return

    if filepath.suffix.lower() != ".pdf":
        _output({"status": "error", "code": "NOT_PDF", "message": "仅支持 PDF 文件"})
        return

    try:
        from pypdf import PdfReader
    except ImportError:
        _output({"status": "error", "code": "MISSING_DEPENDENCY", "message": "缺少 pypdf 依赖"})
        return

    try:
        reader = PdfReader(str(filepath))
        meta = reader.metadata or {}

        result = {"status": "success", "file": str(filepath)}

        if meta.title:
            result["title"] = meta.title
        if meta.author:
            result["authors"] = meta.author
        if meta.subject:
            result["subject"] = meta.subject

        # 从 XMP 元数据中提取 DOI
        doi = None
        if hasattr(reader, 'xmp_metadata') and reader.xmp_metadata:
            xmp = reader.xmp_metadata
            # DOI 可能在 dc:identifier 或自定义属性中
            if hasattr(xmp, 'dc_identifier') and xmp.dc_identifier:
                for ident in (xmp.dc_identifier if isinstance(xmp.dc_identifier, list) else [xmp.dc_identifier]):
                    if ident and '10.' in str(ident):
                        import re as _re
                        doi_m = _re.search(r'(10\.\d{4,}/[^\s]+)', str(ident))
                        if doi_m:
                            doi = doi_m.group(1)
                            break

        # 从前几页文本中查找 DOI
        if not doi:
            import re as _re
            for page_num in range(min(3, len(reader.pages))):
                page_text = reader.pages[page_num].extract_text() or ""
                doi_m = _re.search(r'(?:DOI|doi)[：:\s]*\s*(10\.\d{4,}/[^\s]+)', page_text)
                if doi_m:
                    doi = doi_m.group(1).rstrip(".")
                    break

        if doi:
            result["doi"] = doi
            # 用 DOI 从 Crossref 补全完整元数据
            crossref_data = resolve_crossref(doi)
            if crossref_data:
                result["crossref"] = crossref_data

        _output(result)

    except Exception as e:
        _output({"status": "error", "code": "PDF_READ_FAILED", "message": str(e)})


def cmd_write_docx(args):
    """Markdown 文件 → 学术格式 .docx"""
    md_path = Path(args.filepath)
    if not md_path.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND",
                 "message": f"文件不存在: {args.filepath}"})
        return

    md_text = None
    for enc in ("utf-8", "utf-8-sig", "gbk", "gb2312", "gb18030"):
        try:
            md_text = md_path.read_text(encoding=enc)
            break
        except (UnicodeDecodeError, LookupError):
            continue
    if md_text is None:
        _output({"status": "error", "code": "ENCODING_ERROR",
                 "message": "文件编码无法识别"})
        return

    output_path = Path(args.output) if args.output else md_path.with_suffix(".docx")
    _output(_write_docx_from_markdown(md_text, output_path))


def cmd_patch_docx(args):
    """在现有 .docx 上打补丁：文本替换 + 脚注插入 + 追加参考文献"""
    try:
        from docx import Document
        from docx.shared import Cm
    except ImportError:
        _output({"status": "error", "code": "MISSING_DEPENDENCY",
                 "message": "缺少 python-docx 依赖"})
        return

    docx_path = Path(args.filepath)
    if not docx_path.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND",
                 "message": f"文件不存在: {args.filepath}"})
        return

    patch_path = Path(args.patch)
    if not patch_path.exists():
        _output({"status": "error", "code": "FILE_NOT_FOUND",
                 "message": f"补丁文件不存在: {args.patch}"})
        return

    try:
        patch_data = json.loads(patch_path.read_text(encoding="utf-8"))
    except Exception as e:
        _output({"status": "error", "code": "PATCH_PARSE_FAILED",
                 "message": f"补丁 JSON 解析失败: {e}"})
        return

    if not isinstance(patch_data, dict):
        _output({"status": "error", "code": "PATCH_PARSE_FAILED",
                 "message": "补丁文件顶层必须是 JSON 对象"})
        return

    patches = patch_data.get("patches", [])
    footnotes_list = patch_data.get("footnotes", [])
    append_refs = patch_data.get("append_references", [])

    if not isinstance(patches, list) or not isinstance(footnotes_list, list) or not isinstance(append_refs, list):
        _output({"status": "error", "code": "PATCH_PARSE_FAILED",
                 "message": "patches / footnotes / append_references 必须是数组"})
        return

    output_path = Path(args.output) if args.output else docx_path.with_name(
        docx_path.stem + "_patched" + docx_path.suffix
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(docx_path))
    stats = {"replaced": 0, "not_found": 0, "footnotes_added": 0, "references_appended": 0}
    warnings: List[str] = []

    ftn_element, max_id = _get_or_create_footnotes_part(doc)
    fn_counter = [max_id]

    for patch in patches:
        find_text = patch.get("find", "")
        replace_text = patch.get("replace", "")
        if not find_text:
            continue
        found = False
        for para in doc.paragraphs:
            if find_text in para.text:
                if _para_replace_text(para, find_text, replace_text):
                    stats["replaced"] += 1
                    found = True
                    break
        if not found:
            stats["not_found"] += 1
            warnings.append(f"未找到替换目标: \"{find_text[:30]}...\"" if len(find_text) > 30 else f"未找到替换目标: \"{find_text}\"")

    for fn_entry in footnotes_list:
        after_text = fn_entry.get("after", "")
        fn_text = fn_entry.get("text", "")
        if not after_text or not fn_text:
            continue
        found = False
        for para in doc.paragraphs:
            if after_text not in para.text:
                continue
            run_el = _find_run_containing(para, after_text)
            if run_el is not None:
                _add_footnote_to_element(ftn_element, fn_counter, run_el, fn_text)
                stats["footnotes_added"] += 1
                found = True
                break
        if not found:
            warnings.append(f"脚注定位失败: \"{after_text[:30]}\"" if len(after_text) > 30 else f"脚注定位失败: \"{after_text}\"")

    if append_refs:
        existing_ref_heading = None
        for para in doc.paragraphs:
            if para.text.strip() in ("参考文献", "References") and para.style.name.startswith("Heading"):
                existing_ref_heading = para
        if existing_ref_heading is None:
            doc.add_paragraph()
            try:
                ref_heading = doc.add_paragraph("参考文献")
                ref_heading.style = doc.styles["Heading 1"]
            except KeyError:
                ref_heading = doc.add_paragraph("参考文献")
                ref_heading.runs[0].bold = True
        for ref_text in append_refs:
            p = doc.add_paragraph(ref_text)
            p.paragraph_format.first_line_indent = Cm(-0.74)
            p.paragraph_format.left_indent = Cm(0.74)
        stats["references_appended"] = len(append_refs)

    try:
        doc.save(str(output_path))
    except Exception as e:
        _output({"status": "error", "code": "IO_ERROR",
                 "message": f"保存失败: {e}"})
        return

    has_issues = stats["not_found"] > 0 or len(warnings) > 0
    result: Dict[str, Any] = {
        "status": "partial" if has_issues else "success",
        "message": f"已保存: {output_path}",
        "output": str(output_path),
        **stats,
    }
    if warnings:
        result["warnings"] = warnings
    _output(result)


def add_parser(sub):
    # read-paper
    p_paper = sub.add_parser("read-paper", help="读取论文文件（.docx/.txt/.md）并输出 UTF-8 文本")
    p_paper.add_argument("filepath", help="论文文件路径")
    p_paper.add_argument("--output", help="输出到文件（默认直接打印）")
    p_paper.add_argument("--raw", action="store_true", help="输出纯文本而非 JSON")
    p_paper.set_defaults(func=cmd_read_paper)

    # pdf-meta
    p_pdf = sub.add_parser("pdf-meta", help="从 PDF 提取元数据（标题、DOI 等）")
    p_pdf.add_argument("filepath", help="PDF 文件路径")
    p_pdf.set_defaults(func=cmd_pdf_meta)

    # write-docx
    p_wdocx = sub.add_parser("write-docx", help="Markdown → 学术格式 Word 文档")
    p_wdocx.add_argument("filepath", help="Markdown 文件路径")
    p_wdocx.add_argument("--output", help="输出 .docx 路径（默认同名 .docx）")
    p_wdocx.set_defaults(func=cmd_write_docx)

    # patch-docx
    p_pdocx = sub.add_parser("patch-docx", help="在现有 .docx 上打补丁（插入引用/脚注）")
    p_pdocx.add_argument("filepath", help="原始 .docx 文件路径")
    p_pdocx.add_argument("--patch", required=True, help="补丁 JSON 文件路径")
    p_pdocx.add_argument("--output", help="输出路径（默认 原名_patched.docx）")
    p_pdocx.set_defaults(func=cmd_patch_docx)
