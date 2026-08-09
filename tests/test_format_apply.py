import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

pytest.importorskip("docx")

from core.thesis_format import profile as P  # noqa: E402
from core.thesis_format import apply as A  # noqa: E402
from core.thesis_format import inspect as I  # noqa: E402


def _make_docx(path):
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(1.0)  # 故意不合规
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "仿宋")
    doc.add_heading("摘要", level=1)
    doc.add_paragraph("正文。")
    doc.save(str(path))


def test_apply_to_docx_fixes_styles(tmp_path):
    src = tmp_path / "src.docx"
    out = tmp_path / "out.docx"
    _make_docx(src)
    prof = P.build_template()
    res = A.apply_to_docx(str(src), prof, str(out))
    assert res["status"] == "success"
    assert out.exists()
    m = I.inspect_docx(str(out))
    assert m.body_font_cjk == "宋体"
    assert m.body_font_latin == "Times New Roman"
    assert m.body_size_pt == 12.0
    assert m.page_margin_cm["left"] == 3.17


def test_apply_from_markdown_uses_profile(tmp_path):
    out = tmp_path / "gen.docx"
    prof = P.build_template()
    prof["body"]["font_cjk"] = "楷体"
    res = A.apply_from_markdown("# 摘要\n正文内容。\n", prof, str(out))
    assert res["status"] == "success"
    m = I.inspect_docx(str(out))
    assert m.body_font_cjk == "楷体"
