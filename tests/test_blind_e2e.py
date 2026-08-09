import json
import subprocess
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
ENTRY = SCRIPTS / "literature.py"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

pytest.importorskip("docx")


def _run(*args):
    r = subprocess.run([sys.executable, str(ENTRY), *args], capture_output=True, text=True, encoding="utf-8")
    assert r.returncode == 0, r.stderr
    return json.loads(r.stdout)


def test_review_signals_end_to_end(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("摘要", level=1); doc.add_paragraph("本文研究国际传播。")
    doc.add_heading("第一章 绪论", level=1)
    doc.add_paragraph("本文创新点在于提出新分析框架，基于场域理论。采用问卷调查与内容分析方法。")
    doc.add_heading("第二章 研究现状综述", level=1)
    doc.add_paragraph("已有研究[1][2]表明……" * 20)
    doc.add_heading("第三章 研究设计", level=1); doc.add_paragraph("研究设计内容。")
    doc.add_heading("结论", level=1); doc.add_paragraph("结论内容。")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[1] 张三. 测试[J]. 学报, 2024.")
    doc.add_paragraph("[2] 某. 撤稿研究[J]. 学报, 2020.")
    src = tmp_path / "thesis.docx"; doc.save(str(src))

    rep = _run("review-signals", str(src))
    assert rep["status"] == "success"
    assert set(rep["signals"].keys()) == {"选题与综述", "创新性及论文价值", "基础理论与科研能力", "学术规范与写作水平"}
    assert rep["signals"]["选题与综述"]["measurable"]["lit_review_present"] is True
    assert rep["signals"]["创新性及论文价值"]["measurable"]["innovation_statement_found"] is True
    assert "问卷" in rep["signals"]["基础理论与科研能力"]["measurable"]["method_keywords"]
    assert rep["integrity_flags"]  # 撤稿引用被标记


def test_review_signals_with_format_profile(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_heading("摘要", level=1); doc.add_paragraph("摘要。")
    doc.add_heading("参考文献", level=1); doc.add_paragraph("[1] 张三. 测试[J]. 学报, 2024.")
    src = tmp_path / "t.docx"; doc.save(str(src))

    prof = _run("format-profile", "--template")
    ppath = tmp_path / "p.json"
    ppath.write_text(json.dumps(prof, ensure_ascii=False), encoding="utf-8")

    rep = _run("review-signals", str(src), "--format-profile", str(ppath))
    fc = rep["signals"]["学术规范与写作水平"]["measurable"]["format_check"]
    assert fc is not None and "error" in fc
