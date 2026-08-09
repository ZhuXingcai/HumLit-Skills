import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))


def test_inspect_markdown_extracts_structure():
    from core.thesis_format import inspect as I
    md = (
        "# 摘要\n本文研究[1]。\n\n"
        "# 目录\n第一章 引言\n\n"
        "# 第一章 引言\n正文[1][2]。\n\n"
        "图1-1 框架图\n\n"
        "## 1.1 背景\n内容。\n\n"
        "# 参考文献\n[1] 张三. 测试[J]. 学报, 2024.\n[2] 李四. 再测[J]. 学报, 2025.\n"
    )
    m = I.inspect_markdown(md)
    titles = m.section_titles
    assert "摘要" in titles
    assert any("参考文献" not in t and "引言" in t for t in titles)
    assert len(m.references) == 2
    assert set(m.intext_ref_numbers) >= {1, 2}
    assert any(c.text.startswith("图1-1") for c in m.figure_captions)
    assert "第一章 引言" in m.toc_entries


def test_inspect_docx_reads_styles(tmp_path):
    docx = pytest.importorskip("docx")
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17); sec.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    doc.add_heading("摘要", level=1)
    doc.add_paragraph("正文内容[1]。")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[1] 张三. 测试[J]. 学报, 2024.")
    fp = tmp_path / "t.docx"
    doc.save(str(fp))

    from core.thesis_format import inspect as I
    m = I.inspect_docx(str(fp))
    assert m.body_font_cjk == "宋体"
    assert m.body_font_latin == "Times New Roman"
    assert m.body_size_pt == 12.0
    assert m.page_margin_cm["left"] == 3.17
    assert "摘要" in m.section_titles
    assert len(m.references) == 1
    assert 1 in m.intext_ref_numbers
