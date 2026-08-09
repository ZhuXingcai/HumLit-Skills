import json
import subprocess
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
ENTRY = SCRIPTS / "literature.py"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

pytest.importorskip("docx")


def _run(*args):
    r = subprocess.run([sys.executable, str(ENTRY), *args],
                       capture_output=True, text=True, encoding="utf-8")
    assert r.returncode == 0, r.stderr
    return json.loads(r.stdout)


def test_apply_then_check_closes_loop(tmp_path):
    # 1) 造一份不合规 docx
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    doc = Document()
    doc.sections[0].left_margin = Cm(1.0)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "仿宋")
    doc.add_heading("摘要", level=1); doc.add_paragraph("正文。")
    doc.add_heading("Abstract", level=1); doc.add_paragraph("body.")
    doc.add_heading("目录", level=1); doc.add_paragraph("摘要")
    doc.add_heading("第一章 引言", level=1); doc.add_paragraph("内容[1]。")
    doc.add_heading("参考文献", level=1); doc.add_paragraph("[1] 张三. 测试[J]. 学报, 2024.")
    src = tmp_path / "thesis.docx"; doc.save(str(src))

    # 2) 生成并校验 profile
    prof_path = tmp_path / "profile.json"
    tpl = _run("format-profile", "--template")
    prof_path.write_text(json.dumps(tpl, ensure_ascii=False), encoding="utf-8")
    v = _run("format-profile", "--validate", str(prof_path))
    assert v["status"] == "success"

    # 3) 检测：应发现排版类 error
    before = _run("format-check", str(src), "--profile", str(prof_path))
    assert before["status"] == "success"
    assert before["summary"]["error"] >= 1
    codes = {i["code"] for i in before["issues"]}
    assert "FONT_CJK_MISMATCH" in codes or "MARGIN_MISMATCH" in codes

    # 4) 套用
    out = tmp_path / "thesis_formatted.docx"
    applied = _run("format-apply", str(src), "--profile", str(prof_path),
                   "--output", str(out))
    assert applied["status"] == "success" and out.exists()

    # 5) 再检测：可修的排版 error 应清零
    after = _run("format-check", str(out), "--profile", str(prof_path))
    fixable_errors = [i for i in after["issues"]
                      if i["severity"] == "error" and i["fixable"]]
    assert fixable_errors == [], f"套用后仍有可修 error: {fixable_errors}"
