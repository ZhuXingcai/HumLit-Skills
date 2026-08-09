import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))

from core.thesis_format.inspect import inspect_markdown, inspect_docx  # noqa: E402


def test_markdown_detects_plain_chinese_section_lines():
    text = """摘要
这是摘要。

关键词：数字人文；方志；史料

一、引言
正文引用[1]。

参考文献
[1] 张三. 题名[J]. 期刊, 2020.
"""
    m = inspect_markdown(text)
    assert "摘要" in m.section_titles
    assert "一、引言" in m.section_titles
    assert "参考文献" in m.section_titles
    assert len(m.references) == 1
    assert m.intext_ref_numbers == [1]


def test_markdown_detects_toc_plain_line():
    text = """目录
一、引言
二、方法

一、引言
正文。
"""
    m = inspect_markdown(text)
    assert "目录" in m.section_titles
    assert "一、引言" in m.toc_entries
    assert "二、方法" in m.toc_entries


def test_markdown_skips_blank_lines_inside_toc_before_entries():
    text = """目录

一、引言
二、方法

一、引言
正文。
"""
    m = inspect_markdown(text)
    assert m.toc_entries == ["一、引言", "二、方法"]
    assert "一、引言" in m.section_titles
    assert any(h.text == "一、引言" for h in m.headings)


def test_markdown_exits_toc_when_body_heading_repeats_without_blank():
    text = """目录
一、引言
二、方法
一、引言
正文。
"""
    m = inspect_markdown(text)
    assert m.toc_entries == ["一、引言", "二、方法"]
    assert "一、引言" in m.section_titles
    assert any(h.text == "一、引言" for h in m.headings)
    assert "正文。" not in m.toc_entries


def test_docx_detects_plain_chinese_section_lines(tmp_path):
    from docx import Document

    path = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("摘要")
    doc.add_paragraph("这是摘要。")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("正文引用[1]。")
    doc.add_paragraph("参考文献")
    doc.add_paragraph("[1] 张三. 题名[J]. 期刊, 2020.")
    doc.save(path)

    m = inspect_docx(str(path))
    assert "摘要" in m.section_titles
    assert "一、引言" in m.section_titles
    assert "参考文献" in m.section_titles
    assert len(m.references) == 1


def test_docx_keeps_plain_toc_entries_before_body_heading(tmp_path):
    from docx import Document

    path = tmp_path / "toc.docx"
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("二、方法")
    doc.add_paragraph("")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("正文。")
    doc.save(path)

    m = inspect_docx(str(path))
    assert m.toc_entries == ["一、引言", "二、方法"]
    assert "一、引言" in m.section_titles
    assert any(h.text == "一、引言" for h in m.headings)


def test_docx_exits_toc_when_body_heading_repeats_without_blank(tmp_path):
    from docx import Document

    path = tmp_path / "toc-no-blank.docx"
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("二、方法")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("正文。")
    doc.save(path)

    m = inspect_docx(str(path))
    assert m.toc_entries == ["一、引言", "二、方法"]
    assert "一、引言" in m.section_titles
    assert any(h.text == "一、引言" for h in m.headings)
    assert "正文。" not in m.toc_entries


def test_docx_skips_blank_paragraphs_inside_toc_before_entries(tmp_path):
    from docx import Document

    path = tmp_path / "toc-blank.docx"
    doc = Document()
    doc.add_paragraph("目录")
    doc.add_paragraph("")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("二、方法")
    doc.add_paragraph("1.1 研究背景")
    doc.add_paragraph("")
    doc.add_paragraph("一、引言")
    doc.add_paragraph("正文。")
    doc.save(path)

    m = inspect_docx(str(path))
    assert m.toc_entries == ["一、引言", "二、方法", "1.1 研究背景"]
    assert "一、引言" in m.section_titles
    assert any(h.text == "一、引言" for h in m.headings)
