import json
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
ENTRY = SCRIPTS / "literature.py"
if str(SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SCRIPTS))


def _run(*args):
    r = subprocess.run([sys.executable, str(ENTRY), *args], capture_output=True, text=True, encoding="utf-8")
    assert r.returncode == 0, r.stderr
    return json.loads(r.stdout)


def test_markdown_format_check_marks_page_and_body_unknown(tmp_path):
    prof = tmp_path / "p.json"
    _run("format-profile", "--template", "--output", str(prof))
    draft = tmp_path / "draft.md"
    draft.write_text(
        "# 摘要\n内容\n\n# Abstract\ntext\n\n# 目录\n一、引言\n\n# 一、引言\n正文。\n\n# 参考文献\n[1] 张三. 题名[J]. 期刊, 2020.",
        encoding="utf-8",
    )

    out = _run("format-check", str(draft), "--profile", str(prof))
    assert "page" not in out["summary"]["passed_dimensions"]
    assert "body" not in out["summary"]["passed_dimensions"]
    assert "page" in out["summary"]["unknown_dimensions"]
    assert "body" in out["summary"]["unknown_dimensions"]


def test_docx_with_styles_observes_page_and_body(tmp_path):
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.oxml.ns import qn

    prof = tmp_path / "p.json"
    _run("format-profile", "--template", "--output", str(prof))

    path = tmp_path / "styled.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    doc.add_heading("摘要", level=1)
    doc.add_paragraph("内容")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("text")
    doc.add_heading("目录", level=1)
    doc.add_paragraph("一、引言")
    doc.add_heading("一、引言", level=1)
    doc.add_paragraph("正文。")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[1] 张三. 题名[J]. 期刊, 2020.")
    doc.save(path)

    out = _run("format-check", str(path), "--profile", str(prof))
    assert "page" in out["summary"]["passed_dimensions"]
    assert "body" in out["summary"]["passed_dimensions"]
    assert "page" not in out["summary"]["unknown_dimensions"]
    assert "body" not in out["summary"]["unknown_dimensions"]


def test_docx_with_only_line_spacing_does_not_pass_body(tmp_path):
    from docx import Document
    from docx.shared import Cm

    prof = tmp_path / "p.json"
    _run("format-profile", "--template", "--output", str(prof))

    path = tmp_path / "line-spacing-only.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.5
    doc.add_heading("摘要", level=1)
    doc.add_paragraph("内容")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("text")
    doc.add_heading("目录", level=1)
    doc.add_paragraph("一、引言")
    doc.add_heading("一、引言", level=1)
    doc.add_paragraph("正文。")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[1] 张三. 题名[J]. 期刊, 2020.")
    doc.save(path)

    out = _run("format-check", str(path), "--profile", str(prof))
    assert "body" not in out["summary"]["passed_dimensions"]
    assert "body" in out["summary"]["unknown_dimensions"]


def test_docx_with_letter_paper_does_not_pass_page(tmp_path):
    from docx import Document
    from docx.shared import Cm, Inches, Pt
    from docx.oxml.ns import qn

    prof = tmp_path / "p.json"
    _run("format-profile", "--template", "--output", str(prof))

    path = tmp_path / "letter.docx"
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    doc.add_heading("摘要", level=1)
    doc.add_paragraph("内容")
    doc.add_heading("Abstract", level=1)
    doc.add_paragraph("text")
    doc.add_heading("目录", level=1)
    doc.add_paragraph("一、引言")
    doc.add_heading("一、引言", level=1)
    doc.add_paragraph("正文。")
    doc.add_heading("参考文献", level=1)
    doc.add_paragraph("[1] 张三. 题名[J]. 期刊, 2020.")
    doc.save(path)

    out = _run("format-check", str(path), "--profile", str(prof))
    assert "page" not in out["summary"]["passed_dimensions"]
    assert any(i["code"] == "PAPER_MISMATCH" for i in out["issues"])
